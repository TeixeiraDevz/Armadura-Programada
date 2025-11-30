#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para gerar Especificacao_Requisitos.docx
Requisitos: pip install python-docx
Execute: python gerar_docx.py
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERRO: Biblioteca python-docx não instalada.")
    print("Instale com: pip install python-docx")
    exit(1)

# Criar documento
doc = Document()

# Configurar estilo padrão
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# TÍTULO PRINCIPAL
title = doc.add_heading('ESPECIFICAÇÃO DE REQUISITOS DE SOFTWARE (ERS)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(16)
title_run.font.bold = True

# Adicionar informações do projeto
doc.add_paragraph('Projeto: Armadura Programada')
doc.add_paragraph('Versão do Documento: 1.0')
doc.add_paragraph('Data: 2025')
doc.add_paragraph('')

# 1. INTRODUÇÃO
doc.add_heading('1. INTRODUÇÃO', 1)

doc.add_heading('1.1 Objetivo do Documento', 2)
doc.add_paragraph(
    'Este documento apresenta a Especificação de Requisitos de Software (ERS) '
    'do projeto Armadura Programada, uma plataforma educacional completa para '
    'iniciantes em programação. O documento descreve os requisitos funcionais '
    'e não funcionais do sistema, seus subsistemas, modelos de dados e casos de uso.'
)

doc.add_heading('1.2 Escopo do Sistema', 2)
doc.add_paragraph(
    'O sistema Armadura Programada é uma plataforma web que oferece trilhas de '
    'aprendizado guiadas, projetos práticos, sistema de certificados e uma '
    'comunidade ativa para apoiar o aprendizado de programação do zero ao nível profissional.'
)

doc.add_heading('1.3 Definições e Abreviações', 2)
definitions = [
    ('ERS', 'Especificação de Requisitos de Software'),
    ('PP', 'Plano de Projeto'),
    ('DR', 'Documento de Requisitos'),
    ('IC', 'Item de Configuração'),
    ('API', 'Application Programming Interface'),
    ('JWT', 'JSON Web Token'),
    ('RBAC', 'Role-Based Access Control'),
    ('LGPD', 'Lei Geral de Proteção de Dados')
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Abreviação'
hdr_cells[1].text = 'Definição'
for abbr, definition in definitions:
    row_cells = table.add_row().cells
    row_cells[0].text = abbr
    row_cells[1].text = definition

# 2. DESCRIÇÃO GERAL
doc.add_heading('2. DESCRIÇÃO GERAL', 1)

doc.add_heading('2.1 Perspectiva do Produto', 2)
doc.add_paragraph(
    'O Armadura Programada é um sistema web independente que se integra com '
    'serviços externos para autenticação, armazenamento de arquivos e análise de dados. '
    'O sistema é composto por uma aplicação frontend (React) e uma API backend (Node.js).'
)

doc.add_heading('2.2 Funções do Produto', 2)
doc.add_paragraph('As principais funções do sistema incluem:')
functions = [
    'Gerenciamento de usuários e autenticação',
    'Cadastro e gerenciamento de alunos',
    'Visualização de conteúdo educacional (cursos, trilhas, materiais)',
    'Sistema de avaliação e progresso',
    'Ranking e pontuação para gamificação',
    'Sistema de notificações',
    'Fórum de dúvidas e comunidade',
    'Sistema de recompensas',
    'Integração com ferramentas externas (IDE, GitHub)',
    'Página de perfil do usuário',
    'Blog de notícias e tutoriais'
]
for func in functions:
    p = doc.add_paragraph(func, style='List Bullet')

doc.add_heading('2.3 Características dos Usuários', 2)
doc.add_paragraph(
    'O sistema atende principalmente a iniciantes em programação, mas também '
    'suporta usuários avançados que desejam aprimorar seus conhecimentos.'
)

doc.add_heading('2.4 Restrições Gerais', 2)
restrictions = [
    'O sistema deve ser compatível com navegadores modernos (Chrome, Firefox, Safari, Edge)',
    'Requisitos de segurança conforme LGPD',
    'Disponibilidade 24/7 com tempo de resposta adequado',
    'Suporte para dispositivos móveis e desktop'
]
for restriction in restrictions:
    p = doc.add_paragraph(restriction, style='List Bullet')

# 3. ESPECIFICAÇÃO DE REQUISITOS
doc.add_heading('3. ESPECIFICAÇÃO DE REQUISITOS', 1)

doc.add_heading('3.1 Identificação de Subsistemas', 2)
doc.add_paragraph(
    'O sistema é composto pelos seguintes subsistemas principais:'
)

subsystems = [
    {
        'nome': 'Gerenciamento de Usuários',
        'descricao': 'Responsável pela autenticação e controle de autorização dentro do sistema. Inclui registro de usuários, gerenciamento de papéis, controle de permissões (RBAC) e registro de log de atividades.'
    },
    {
        'nome': 'Cadastro de Clientes',
        'descricao': 'Gerencia o cadastro e manutenção de dados de clientes, incluindo informações industriais, financeiras e de conformidade com LGPD. Integra-se com o módulo de documentos e pipelines de análise.'
    },
    {
        'nome': 'Módulo de IA (Análise e Score)',
        'descricao': 'Fornece pipelines de análise de crédito e scoring, pré-processamento de dados e classificação de risco baseada em modelos de machine learning. Permite também treinamento e verificação de novos modelos.'
    },
    {
        'nome': 'Relatórios e Métricas',
        'descricao': 'Responsável pela geração de relatórios e dashboards para usuários. Inclui relatórios de aprovação/reprovação, exportações PORCSY e indicadores operacionais e preditivos.'
    },
    {
        'nome': 'Gerenciamento de Documentos',
        'descricao': 'Subsistema responsável pelo upload, validação e armazenamento de documentos financeiros e outros. Garante integridade dos arquivos e integra-se com o módulo de IA.'
    },
    {
        'nome': 'Auditoria e Conformidade',
        'descricao': 'Controla logs de auditoria, eventos do sistema e logs de sistema, garantindo rastreabilidade e conformidade com políticas de segurança e LGPD.'
    },
    {
        'nome': 'Infraestrutura e Configuração',
        'descricao': 'Responsável pelo controle de ambientes (dev, stage, prod), deployments e gerenciamento de configuração. Trabalha em conjunto com Configuração e DBA.'
    }
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Subsistema'
hdr_cells[1].text = 'Descrição'
for sub in subsystems:
    row_cells = table.add_row().cells
    row_cells[0].text = sub['nome']
    row_cells[1].text = sub['descricao']

doc.add_paragraph('')
doc.add_paragraph('Figura 1 - Diagrama de Entidades', style='Caption')
doc.add_paragraph(
    '[Nota: O diagrama de entidades mostra a arquitetura de alto nível do sistema, '
    'com a API central conectando os módulos WEB, CLIENTS, REPORTS e AUDIT. '
    'Os módulos USERS, DOCS e IA estão conectados ao módulo CLIENTS.]'
)

# 3.2 Modelo de Caso de Uso
doc.add_heading('3.2 Modelo de Caso de Uso', 2)

doc.add_heading('3.2.1 Atores do Sistema', 3)

actors = [
    {
        'nome': 'Administrador',
        'descricao': 'Usuário responsável pela gestão do sistema, permissões e monitoramento das análises.',
        'funcoes': [
            'Gerenciar contas de usuários e permissões',
            'Visualizar métricas e relatórios globais',
            'Supervisionar o modelo de IA'
        ]
    },
    {
        'nome': 'Analista de Crédito',
        'descricao': 'Profissional da instituição financeira responsável por realizar análises e consultar relatórios.',
        'funcoes': [
            'Cadastrar solicitantes',
            'Fazer upload de dados financeiros',
            'Executar análises e visualizar relatórios'
        ]
    },
    {
        'nome': 'Solicitante (Indireto)',
        'descricao': 'Indivíduo ou empresa avaliado pelo sistema.',
        'funcoes': [
            'Não acessa o sistema diretamente',
            'Seus dados são processados pelos analistas para geração de score'
        ]
    }
]

for actor in actors:
    doc.add_heading(f'Ator: {actor["nome"]}', 4)
    doc.add_paragraph(actor['descricao'])
    doc.add_paragraph('Funções Principais:')
    for func in actor['funcoes']:
        p = doc.add_paragraph(func, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('Figura 2.3.1 - Descrição de Funções específicas de cada ator', style='Caption')

# 3.3 Diagramas de Classes
doc.add_heading('3.3 Diagramas de Classes', 2)

doc.add_paragraph('Figura 10: Diagrama de Classes sem atributos')
doc.add_paragraph(
    '[Nota: O diagrama mostra as classes principais: Usuario, Aluno, Professor, '
    'Administrador, Transacao, Pagamento, Relatorio, Curso e Disciplina, com seus relacionamentos.]'
)

doc.add_paragraph('')
doc.add_paragraph('Figura 11: Diagrama de Classes com atributos')
doc.add_paragraph(
    '[Nota: O diagrama detalhado inclui atributos e multiplicidades das relações. '
    'Usuario é a classe base, com Aluno, Professor e Administrador como subclasses. '
    'Relacionamentos incluem: Usuario (1) -> Transacao (0..*), Usuario (1) -> Pagamento (0..*), '
    'Professor (1) -> Curso (0..*), Curso (1) -> Disciplina (0..*).]'
)

# 3.4 Modelo Estrutural
doc.add_heading('3.4 Modelo Estrutural', 2)

doc.add_paragraph('Figura 13 - Modelo Estrutural')
doc.add_paragraph(
    'O modelo estrutural define as seguintes entidades principais:'
)

entities = [
    {
        'nome': 'CLIENTE',
        'atributos': 'id, nome, cpfCnpj, estabelecimento, inscricaoEstadual, endereco, comandoGeral, dataGeracaoCredito, criadoEm, atualizadoEm'
    },
    {
        'nome': 'USUÁRIO',
        'atributos': 'id, nome, email, senhaHash, papel, criadoEm, atualizadoEm'
    },
    {
        'nome': 'DOCUMENTO',
        'atributos': 'id, id_cliente (FK), tipo, nomeArquivo, caminhoArquivo, dataUpload, validado'
    },
    {
        'nome': 'ANÁLISE',
        'atributos': 'id, id_cliente (FK), id_usuario (FK), id_documento (FK), pontuacao, risco, criadoEm, vinculoSocial, analiseDe, observacoes'
    },
    {
        'nome': 'RELATÓRIO',
        'atributos': 'id, id_analise (FK), tipo, caminhoArquivo, dataGeracao'
    }
]

for entity in entities:
    doc.add_paragraph(f'{entity["nome"]}: {entity["atributos"]}', style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('Relacionamentos:')
relationships = [
    'CLIENTE possui DOCUMENTO (1:N)',
    'CLIENTE é referenciado em ANÁLISE (1:N)',
    'USUÁRIO realiza ANÁLISE (1:N)',
    'DOCUMENTO é analisado em ANÁLISE (1:N)',
    'ANÁLISE gera RELATÓRIO (1:N)'
]
for rel in relationships:
    p = doc.add_paragraph(rel, style='List Bullet')

# 3.5 Tabela de Necessidades
doc.add_heading('3.5 Tabela de Necessidades', 2)

needs = [
    {
        'codigo': 'N001',
        'nome': 'Plataforma Modular',
        'descricao': 'Plataforma administrável com lógica visível (módulos, pacotes e gerenciamento).',
        'status': 'Inexistente',
        'resultado': 'Nova estrutura modular e lógica legal.'
    },
    {
        'codigo': 'N002',
        'nome': 'Login Seguro',
        'descricao': 'Sistema de autenticação seguro com estratégia de autorização e proteção de dados (JWT).',
        'status': 'Inexistente',
        'resultado': 'Sistema de autenticação seguro e robusto.'
    },
    {
        'codigo': 'N003',
        'nome': 'Plataforma Escalável',
        'descricao': 'Preparação para futuras expansões da plataforma e otimização de cultura.',
        'status': 'Inexistente',
        'resultado': 'Sistema escalável preparado para cultura.'
    },
    {
        'codigo': 'N004',
        'nome': 'Registro de Usuário',
        'descricao': 'Sistema de registro (nome, email, senha) e validação de dados.',
        'status': 'Inexistente',
        'resultado': 'Sistema de registro e proteção de usuários.'
    },
    {
        'codigo': 'N005',
        'nome': 'Plataforma Multiplataforma',
        'descricao': 'Plataforma acessível em diferentes dispositivos e sistemas operacionais.',
        'status': 'Inexistente',
        'resultado': 'Sistema multiplataforma pronto.'
    },
    {
        'codigo': 'N006',
        'nome': 'Visualização de Conteúdo',
        'descricao': 'Sistema para visualização de conteúdo em aulas, cursos e organização de materiais.',
        'status': 'Inexistente',
        'resultado': 'Sistema de visualização e organização de materiais.'
    },
    {
        'codigo': 'N007',
        'nome': 'Sistema de Avaliação',
        'descricao': 'Sistema de avaliação de progresso, com testes e desafios.',
        'status': 'Inexistente',
        'resultado': 'Sistema de avaliação e acompanhamento de progresso.'
    },
    {
        'codigo': 'N008',
        'nome': 'Ranking e Pontuação',
        'descricao': 'Sistema de ranking e pontuação para motivar o aprendizado.',
        'status': 'Inexistente',
        'resultado': 'Sistema de ranking e motivação.'
    },
    {
        'codigo': 'N009',
        'nome': 'Notificações',
        'descricao': 'Sistema de notificações para alertas, atualizações e lembretes.',
        'status': 'Inexistente',
        'resultado': 'Sistema de notificações e lembretes.'
    },
    {
        'codigo': 'N010',
        'nome': 'Fórum de Dúvidas',
        'descricao': 'Fórum para interação entre alunos e professores, com suporte.',
        'status': 'Inexistente',
        'resultado': 'Fórum de suporte e interação.'
    },
    {
        'codigo': 'N011',
        'nome': 'Sistema de Recompensas',
        'descricao': 'Sistema de recompensas para incentivar aprendizado e participação.',
        'status': 'Inexistente',
        'resultado': 'Sistema de recompensas e incentivos.'
    },
    {
        'codigo': 'N012',
        'nome': 'Integração com Ferramentas',
        'descricao': 'Integração com ferramentas externas (IDE, GitHub, etc.) para projetos.',
        'status': 'Inexistente',
        'resultado': 'Integração com ferramentas e projetos.'
    },
    {
        'codigo': 'N013',
        'nome': 'Página de Perfil',
        'descricao': 'Página de perfil do usuário com informações, progresso e conquistas.',
        'status': 'Inexistente',
        'resultado': 'Página de perfil e conquistas.'
    },
    {
        'codigo': 'N014',
        'nome': 'Blog de Notícias',
        'descricao': 'Blog com notícias, artigos e tutoriais sobre programação.',
        'status': 'Inexistente',
        'resultado': 'Blog de notícias e tutoriais.'
    }
]

table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Código'
hdr_cells[1].text = 'Nome da Necessidade'
hdr_cells[2].text = 'Descrição da Necessidade'
hdr_cells[3].text = 'Status Atual'
hdr_cells[4].text = 'Resultado Esperado'

for need in needs:
    row_cells = table.add_row().cells
    row_cells[0].text = need['codigo']
    row_cells[1].text = need['nome']
    row_cells[2].text = need['descricao']
    row_cells[3].text = need['status']
    row_cells[4].text = need['resultado']

# 4. REQUISITOS NÃO FUNCIONAIS
doc.add_heading('4. REQUISITOS NÃO FUNCIONAIS', 1)

doc.add_heading('4.1 Performance', 2)
doc.add_paragraph(
    'O sistema deve responder a requisições em tempo adequado, com tempo de '
    'resposta médio inferior a 2 segundos para operações comuns e suporte para '
    'múltiplos usuários simultâneos.'
)

doc.add_heading('4.2 Segurança', 2)
doc.add_paragraph(
    'O sistema deve implementar autenticação segura (JWT), criptografia de dados '
    'sensíveis, controle de acesso baseado em papéis (RBAC) e conformidade com LGPD.'
)

doc.add_heading('4.3 Escalabilidade', 2)
doc.add_paragraph(
    'A arquitetura deve suportar crescimento horizontal, permitindo adicionar novos '
    'servidores conforme a demanda aumenta.'
)

doc.add_heading('4.4 Usabilidade', 2)
doc.add_paragraph(
    'A interface deve ser intuitiva, responsiva e acessível, seguindo padrões de '
    'acessibilidade web (WCAG 2.1).'
)

# 5. GLOSSÁRIO
doc.add_heading('5. GLOSSÁRIO', 1)

glossary_terms = [
    ('API', 'Interface de Programação de Aplicações - conjunto de rotinas e padrões para acesso a um aplicativo'),
    ('JWT', 'JSON Web Token - padrão para autenticação e autorização'),
    ('RBAC', 'Role-Based Access Control - controle de acesso baseado em papéis'),
    ('LGPD', 'Lei Geral de Proteção de Dados - legislação brasileira sobre proteção de dados'),
    ('ERD', 'Entity-Relationship Diagram - diagrama de relacionamento entre entidades'),
    ('SCM', 'Software Configuration Management - gerenciamento de configuração de software')
]

for term, definition in glossary_terms:
    p = doc.add_paragraph()
    p.add_run(f'{term}: ').bold = True
    p.add_run(definition)

# Salvar documento
output_file = 'Especificacao_Requisitos.docx'
doc.save(output_file)
print(f"✅ Documento {output_file} gerado com sucesso!")
print(f"📄 Arquivo salvo em: {output_file}")

